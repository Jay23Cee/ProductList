import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pandas as pd
import os

TEMPLATE_FILE = 'template_productlist.docx'


def add_product(doc, product):
    # Add a page break
    doc.add_page_break()

    # Add the non-tabulated fields as paragraphs
    for field, value in product.items():
        if field not in ['Cost', 'Shipping', 'Total cost', 'My Price Range', 'Notes']:  # Exclude 'Notes' from this section
            # Create a new paragraph with the field and value
            p = doc.add_paragraph()
            run = p.add_run(f'{field}: {value}')
            run.font.size = Pt(12)

    # Create a table for tabulated fields
    table = doc.add_table(rows=0, cols=2)

    # Fill the table with the tabulated fields
    tabulated_fields = ['Cost', 'Shipping', 'Total cost', 'My Price Range']
    for i in range(0, len(tabulated_fields), 2):  # step by 2 as we're filling two columns
        cells = table.add_row().cells
        cells[0].text = f'{tabulated_fields[i]}: {str(product[tabulated_fields[i]])}'
        if i+1 < len(tabulated_fields):  # check if there is a field for the second column
            cells[1].text = f'{tabulated_fields[i+1]}: {str(product[tabulated_fields[i+1]])}'

    # Add a space after the table
    doc.add_paragraph()

    # Add the 'Notes' field as a paragraph at the end
    if 'Notes' in product:
        p = doc.add_paragraph()
        run = p.add_run(f'Notes: {product["Notes"]}')
        run.font.size = Pt(12)


root = tk.Tk()

def upload_file():
    file_path = filedialog.askopenfilename(filetypes=[('Excel Files', '*.xlsx'), ('CSV Files', '*.csv')])
    if file_path:
        doc = Document(TEMPLATE_FILE)

        # Read the Excel or CSV file and process each row
        if file_path.endswith('.xlsx'):
            df = pd.read_excel(file_path)
        else:
            df = pd.read_csv(file_path, delimiter='\t')

        # Iterate over each row in the Excel or CSV file
        for index, row in df.iterrows():
            product = {
                "Item Name": row['item_name'],
                "Link": row['link'],
                "Cost": row['cost'],
                "Shipping": row['shipping'],
                "Total cost": row['total'],
                "My Price Range": "$" + str(row['my_price_2x']) + "-" + "$" + str(row['my_price_4x']),
                "Notes": row['notes']
            }

            # Add the product to the document
            add_product(doc, product)

        # Save the Word document
        current_directory = os.getcwd()
        save_path = os.path.join(current_directory, "output.docx")
        doc.save(save_path)

upload_button = tk.Button(root, text="Upload File", command=upload_file)
upload_button.pack()

root.mainloop()
